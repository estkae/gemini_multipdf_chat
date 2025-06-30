import os
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import streamlit as st
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from pathlib import Path
import docx
import glob
import pytesseract
from PIL import Image
import pdf2image
import io
import tempfile

load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# OCR functions
def perform_ocr_on_image(image):
    """Perform OCR on a PIL Image object"""
    try:
        text = pytesseract.image_to_string(image, lang='eng+deu')  # Support English and German
        return text
    except Exception as e:
        st.warning(f"OCR failed: {str(e)}")
        return ""

def extract_text_from_pdf_with_ocr(file_path):
    """Extract text from PDF using OCR (for scanned PDFs)"""
    text = ""
    try:
        # Convert PDF to images
        images = pdf2image.convert_from_path(file_path)
        
        # Perform OCR on each page
        for i, image in enumerate(images):
            page_text = perform_ocr_on_image(image)
            text += f"\n--- Page {i+1} ---\n{page_text}\n"
            
    except Exception as e:
        if "poppler" in str(e).lower():
            st.error("❌ Poppler nicht gefunden! Für OCR-Unterstützung installieren Sie bitte Poppler:")
            st.info("**Windows:** Laden Sie Poppler von https://github.com/oschwartz10612/poppler-windows/releases/ herunter und fügen Sie es zum PATH hinzu")
            st.info("**Oder mit Conda:** conda install -c conda-forge poppler")
        else:
            st.warning(f"PDF OCR failed: {str(e)}")
    
    return text

def read_image_file(file_path):
    """Read text from image files using OCR"""
    try:
        image = Image.open(file_path)
        text = perform_ocr_on_image(image)
        return text
    except Exception as e:
        st.warning(f"Image reading failed: {str(e)}")
        return ""

# read text from various file types
def read_txt_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def read_docx_file(file_path):
    doc = docx.Document(file_path)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

def read_pdf_file(file_path):
    """Read PDF with fallback to OCR if needed"""
    text = ""
    needs_ocr = False
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                
                # Check if page has meaningful text
                if page_text and len(page_text.strip()) > 10:
                    text += page_text
                else:
                    needs_ocr = True
                    break
        
        # If regular extraction failed or returned minimal text, use OCR
        if needs_ocr or len(text.strip()) < 50:
            st.info(f"Using OCR for {os.path.basename(file_path)}...")
            ocr_text = extract_text_from_pdf_with_ocr(file_path)
            if ocr_text:
                text = ocr_text
                
    except Exception as e:
        st.warning(f"PDF reading failed, trying OCR: {str(e)}")
        text = extract_text_from_pdf_with_ocr(file_path)
    
    return text

# read all uploaded files and return text
def get_pdf_text(uploaded_files):
    text = ""
    for file in uploaded_files:
        try:
            file_extension = Path(file.name).suffix.lower()
            
            # Save file temporarily for processing
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                tmp_file.write(file.getbuffer())
                tmp_file_path = tmp_file.name
            
            try:
                if file_extension == '.pdf':
                    # Use the enhanced PDF reader with OCR fallback
                    file_text = read_pdf_file(tmp_file_path)
                    text += file_text
                elif file_extension == '.txt':
                    file.seek(0)  # Reset file pointer
                    text += str(file.read(), "utf-8")
                elif file_extension in ['.docx', '.doc']:
                    doc = docx.Document(tmp_file_path)
                    text += '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                    st.info(f"Using OCR for image {file.name}...")
                    image_text = read_image_file(tmp_file_path)
                    text += image_text
                    
                text += f"\n\n--- End of {file.name} ---\n\n"
            finally:
                # Clean up temp file
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
                    
        except Exception as e:
            st.warning(f"Could not process {file.name}: {str(e)}")
    
    return text

# read files from folder
def get_folder_text(folder_path):
    text = ""
    processed_files = []
    
    # Define supported file extensions
    supported_extensions = {
        '.pdf': read_pdf_file,
        '.txt': read_txt_file,
        '.docx': read_docx_file,
        '.doc': read_docx_file,
        '.png': read_image_file,
        '.jpg': read_image_file,
        '.jpeg': read_image_file,
        '.tiff': read_image_file,
        '.bmp': read_image_file,
    }
    
    # Get all files in the folder
    for ext, reader_func in supported_extensions.items():
        pattern = os.path.join(folder_path, f"*{ext}")
        files = glob.glob(pattern)
        
        for file_path in files:
            try:
                if ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                    st.info(f"Using OCR for image {os.path.basename(file_path)}...")
                file_text = reader_func(file_path)
                text += f"\n\n--- Content from {os.path.basename(file_path)} ---\n\n"
                text += file_text
                processed_files.append(os.path.basename(file_path))
            except Exception as e:
                st.warning(f"Could not read {os.path.basename(file_path)}: {str(e)}")
    
    return text, processed_files

# split text into chunks


def get_text_chunks(text):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=10000, chunk_overlap=1000)
    chunks = splitter.split_text(text)
    return chunks  # list of strings

# get embeddings for each chunk


def get_vector_store(chunks):
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001")  # type: ignore
    vector_store = FAISS.from_texts(chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")


def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context, make sure to provide all the details, if the answer is not in
    provided context just say, "answer is not available in the context", don't provide the wrong answer\n\n
    Context:\n {context}?\n
    Question: \n{question}\n

    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash",
                                   client=genai,
                                   temperature=0.3,
                                   )
    prompt = PromptTemplate(template=prompt_template,
                            input_variables=["context", "question"])
    chain = load_qa_chain(llm=model, chain_type="stuff", prompt=prompt)
    return chain


def clear_chat_history():
    st.session_state.messages = [
        {"role": "assistant", "content": "Upload PDFs or specify a folder path, then ask me a question"}]


def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001")  # type: ignore

    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True) 
    docs = new_db.similarity_search(user_question)

    chain = get_conversational_chain()

    response = chain(
        {"input_documents": docs, "question": user_question}, return_only_outputs=True, )

    print(response)
    return response


def main():
    st.set_page_config(
        page_title="Gemini Document Chatbot",
        page_icon="🤖"
    )

    # Sidebar for uploading PDF files or selecting folder
    with st.sidebar:
        st.title("Menu:")
        
        # Add tabs for different input methods
        tab1, tab2 = st.tabs(["Upload Files", "Select Folder"])
        
        with tab1:
            pdf_docs = st.file_uploader(
                "Upload your files (PDFs, Images, Documents)", accept_multiple_files=True, 
                type=['pdf', 'txt', 'docx', 'doc', 'png', 'jpg', 'jpeg', 'tiff', 'bmp'])
            st.caption("📝 Supports: PDF, TXT, DOCX, DOC, PNG, JPG, JPEG, TIFF, BMP")
            st.caption("🔍 OCR will be used automatically for scanned PDFs and images")
            if st.button("Process Uploaded Files", key="upload_btn"):
                if pdf_docs:
                    with st.spinner("Processing uploaded files..."):
                        raw_text = get_pdf_text(pdf_docs)
                        text_chunks = get_text_chunks(raw_text)
                        get_vector_store(text_chunks)
                        st.success(f"Processed {len(pdf_docs)} files successfully!")
                else:
                    st.error("Please upload at least one file")
        
        with tab2:
            folder_path = st.text_input("Enter folder path:", 
                                       placeholder="/path/to/your/folder")
            if st.button("Process Folder", key="folder_btn"):
                if folder_path and os.path.exists(folder_path):
                    with st.spinner("Processing files from folder..."):
                        raw_text, processed_files = get_folder_text(folder_path)
                        if raw_text:
                            text_chunks = get_text_chunks(raw_text)
                            get_vector_store(text_chunks)
                            st.success(f"Processed {len(processed_files)} files:")
                            for file in processed_files:
                                st.write(f"✓ {file}")
                        else:
                            st.error("No supported files found in the folder")
                elif folder_path:
                    st.error("Folder path does not exist")
                else:
                    st.error("Please enter a folder path")

    # Main content area for displaying chat messages
    st.title("Chat with Documents using Gemini🤖")
    st.write("Welcome to the chat!")
    st.sidebar.button('Clear Chat History', on_click=clear_chat_history)

    # Chat input
    # Placeholder for chat messages

    if "messages" not in st.session_state.keys():
        st.session_state.messages = [
            {"role": "assistant", "content": "Upload PDFs or specify a folder path, then ask me a question"}]

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.write(message["content"])

    if prompt := st.chat_input():
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.write(prompt)

    # Display chat messages and bot response
    if st.session_state.messages[-1]["role"] != "assistant":
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                response = user_input(prompt)
                placeholder = st.empty()
                full_response = ''
                for item in response['output_text']:
                    full_response += item
                    placeholder.markdown(full_response)
                placeholder.markdown(full_response)
        if response is not None:
            message = {"role": "assistant", "content": full_response}
            st.session_state.messages.append(message)


if __name__ == "__main__":
    main()
